#!/usr/bin/python
# encoding:utf-8
"""
Format all the laws from txt and docx to the level of kuan.
"""
import sys
sys.path.append('/data/private/hxy/PLM')
import os
from os.path import join
import traceback
import json

from docx import Document
from cn2an import cn2an, an2cn

from build_corpus.const import *
from build_corpus.config import *
from build_corpus.helper import read_formatted_laws


def _cn_add(ch):
    """
    Add one to the chinese character.

    :param ch:
    :return:
    """

    return an2cn(cn2an(ch, mode='normal') + 1, 'low')


def format_national_xzfg_dfxfg(path, type_, output_path):
    """
    Format {xzfg,dfxfg} parsed from national legal docs database(国家法律法规数据库).

    :param path:
    :param type_:
    :param output_path:
    :return:
    """


    laws = []

    for fp_num, fp in enumerate(os.listdir(path)):
        # get lawname

        # print(fp_num)

        try:

            doc = Document(join(path, fp))

        except:
            print(traceback.print_exc())
            print(fp)
            continue

        chapter_list = []
        cur_chapter = {}
        cur_term_num = ''
        cur_kuan_num = ''

        title = ""
        head = ""

        para_index = 0

        # skip space
        while para_index < len(doc.paragraphs):
            para = doc.paragraphs[para_index].text
            para = para.strip()
            para = para.replace("\u3000", "")
            para = para.replace("\n", "")
            para = para.replace("\xa0", "")
            para = para.replace(" ", "")
            para = para.replace(" ", "")

            if len(para) == 0:
                para_index += 1
            else:
                break

        # get title
        while para_index < len(doc.paragraphs):
            para = doc.paragraphs[para_index].text
            para = para.strip()
            para = para.replace("\u3000", "")
            para = para.replace("\n", "")
            para = para.replace("\xa0", "")
            para = para.replace(" ", "")
            para = para.replace(" ", "")
            para_index += 1

            if len(para) != 0 and not para.startswith("（") and not para.startswith("("):
                title += para
            else:
                break

        # get head
        while para_index < len(doc.paragraphs):
            para = doc.paragraphs[para_index].text
            para = para.strip()
            para = para.replace("\u3000", "")
            para = para.replace("\n", "")
            para = para.replace("\xa0", "")
            para = para.replace(" ", "")
            para = para.replace(" ", "")

            if para.startswith("第一条"):
                break

            head += para

            para_index += 1

        # get chapter
        while para_index < len(doc.paragraphs):
            line = doc.paragraphs[para_index].text
            line = line.strip()
            line = line.replace("\u3000", "")
            line = line.replace("\n", "")
            line = line.replace("\xa0", "")
            line = line.replace(" ", "")
            line = line.replace(" ", "")

            # whether is a chapter
            res = re.search(p_chapter1, line)
            if res is not None:
                if len(cur_chapter) != 0:
                    chapter_list.append(cur_chapter)
                    cur_chapter = {}

                para_index += 1
                continue  # skip chapter

            # whether is a term
            res = re.search(p_term_content1, line)
            if res is not None:
                cur_term_num = res.group("term_num")
                cur_kuan_num = '一'
                cur_kuan = res.group("content")
                cur_term = {cur_kuan_num: cur_kuan}

                cur_chapter[cur_term_num] = cur_term

                # [{'第一条': {'第一款': {{}}, '第二款': {}}}]

            else:

                # 分章节的地方 可能会跨越章节
                # 检查是不是项

                if cur_term_num in cur_chapter.keys():

                    if re.search(p_xiang_content, line) is not None:
                        cur_chapter[cur_term_num][cur_kuan_num] += line
                    else:
                        cur_kuan_num = _cn_add(cur_kuan_num)
                        cur_chapter[cur_term_num][cur_kuan_num] = line

            para_index += 1

        chapter_list.append(cur_chapter)

        law = {'title': title, 'type': type_, 'head': head, 'content': chapter_list, 'file_path': fp}

        laws.append(law)

        if len(title) > 50:
            print(title)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(laws, f, ensure_ascii=False)


def format_pku_fl_sfjs_with_date(law_path, type_, name_date_path, output_path):
    """
    Format {fl,sfjs} downloaded from pku(北大法宝), meanwhile add the implemantion date
    throuth add_fl_sfjs_date.

    :param law_path:
    :param type_:
    :param name_date_path:
    :param output_path:
    :return:
    """


    laws = []

    for fp_num, fp in enumerate(os.listdir(law_path)):

        print(fp_num)
        with open(law_path + '/' + fp, "r", encoding='utf-8') as f:
            doc = [line for line in f.readlines()]
        whole_text = "".join(doc)

        chapter_list = []
        cur_chapter = {}
        cur_term_num = ''
        cur_kuan_num = ''

        title = re.search(p_lawfile, fp).group('lawname')
        print(title)
        head = ""

        para_index = 0

        if '第一条' in whole_text and '第二条' in whole_text:

            # get head
            while para_index < len(doc):
                para = doc[para_index]
                para = para.strip()
                para = para.replace("\u3000", "")
                para = para.replace("\n", "")
                para = para.replace("\xa0", "")
                para = para.replace(" ", "")
                para = para.replace(" ", "")

                if para.startswith("第一条"):
                    break

                head += para

                para_index += 1

            # get chapter
            while para_index < len(doc):
                line = doc[para_index]
                line = line.strip()
                line = line.replace("\u3000", "")
                line = line.replace("\n", "")
                line = line.replace("\xa0", "")
                line = line.replace(" ", "")
                line = line.replace(" ", "")

                # whether is a chapter
                res = re.search(p_chapter1, line)
                if res is not None:
                    if len(cur_chapter) != 0:
                        chapter_list.append(cur_chapter)
                        cur_chapter = {}

                    para_index += 1
                    continue  # skip chapter

                # whether is a term
                res = re.search(p_term_content1, line)
                if res is not None:
                    cur_term_num = res.group("term_num")
                    cur_kuan_num = '一'
                    cur_kuan = res.group("content")
                    cur_term = {cur_kuan_num: cur_kuan}

                    cur_chapter[cur_term_num] = cur_term

                    # [{'第一条': {'第一款': {{}}, '第二款': {}}}]

                else:

                    # 分章节的地方 可能会跨越章节
                    # 检查是不是项

                    if cur_term_num in cur_chapter.keys():

                        if re.search(p_xiang_content, line) is not None:
                            cur_chapter[cur_term_num][cur_kuan_num] += line
                        else:
                            cur_kuan_num = _cn_add(cur_kuan_num)
                            cur_chapter[cur_term_num][cur_kuan_num] = line

                para_index += 1

            chapter_list.append(cur_chapter)




        else:

            # get head
            while para_index < len(doc):
                para = doc[para_index]
                para = para.strip()
                para = para.replace("\u3000", "")
                para = para.replace("\n", "")
                para = para.replace("\xa0", "")
                para = para.replace(" ", "")
                para = para.replace(" ", "")

                if para.startswith("一、"):
                    break

                head += para

                para_index += 1

            # get chapter
            while para_index < len(doc):
                line = doc[para_index]
                line = line.strip()
                line = line.replace("\u3000", "")
                line = line.replace("\n", "")
                line = line.replace("\xa0", "")
                line = line.replace(" ", "")
                line = line.replace(" ", "")

                # # whether is a chapter
                # res = re.search(p_chapter2, line)
                # if res is not None:
                #     if len(cur_chapter) != 0:
                #         chapter_list.append(cur_chapter)
                #         cur_chapter = {}
                #
                #     para_index += 1
                #     continue  # skip chapter

                # whether is a term
                res = re.search(p_term_content2, line)
                if res is not None:
                    cur_term_num = res.group("term_num")
                    cur_kuan_num = '一'
                    cur_kuan = res.group("content")
                    cur_term = {cur_kuan_num: cur_kuan}

                    cur_chapter[cur_term_num] = cur_term

                    # [{'第一条': {'第一款': {{}}, '第二款': {}}}]

                else:

                    # 分章节的地方 可能会跨越章节
                    # 检查是不是项

                    if cur_term_num in cur_chapter.keys():

                        if re.search(p_xiang_content, line) is not None:
                            cur_chapter[cur_term_num][cur_kuan_num] += line
                        else:
                            cur_kuan_num = _cn_add(cur_kuan_num)
                            cur_chapter[cur_term_num][cur_kuan_num] = line

                para_index += 1

            chapter_list.append(cur_chapter)

        law = {'title': title, 'type': type_, 'head': head, 'content': chapter_list, 'file_path': fp}

        laws.append(law)

    add_fl_sfjs_date(laws, name_date_path, output_path)


def add_fl_sfjs_date(formatted_law, name_date_path, output_path):
    """
    Add the implementation date of fl or sfjs from {fl,sfjs}_name_data.json.

    :param formatted_law:
    :param name_date_path:
    :param output_path:
    :return:
    """

    with open(name_date_path, 'r', encoding='utf-8') as f:
        name_dates = json.load(f)

    for fl in formatted_law:
        for name_date in name_dates:
            if name_date['name'] in fl['file_path']:
                fl.update({'date': name_date['date']})
                break
        else:
            print(fl['title'])

    print(formatted_law[:5])

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(formatted_law, f, ensure_ascii=False)


def add_law_ids():
    """
    Index all the formatted laws from zero. It is in-place and the index
    will be saved to the same file.

    :return:
    """

    law_id = 0
    formatted_fl, formatted_sfjs, formatted_xzfg, formatted_dfxfg = read_formatted_laws()

    for i in formatted_fl:
        i["id_"] = law_id
        law_id += 1

    for i in formatted_sfjs:
        i["id_"] = law_id
        law_id += 1

    for i in formatted_xzfg:
        i["id_"] = law_id
        law_id += 1

    for i in formatted_dfxfg:
        i["id_"] = law_id
        law_id += 1

    with open(formatted_fl_path, 'w', encoding='utf-8') as f:
        json.dump(formatted_fl, f, ensure_ascii=False)

    with open(formatted_sfjs_path, 'w', encoding='utf-8') as f:
        json.dump(formatted_sfjs, f, ensure_ascii=False)

    with open(formatted_xzfg_path, 'w', encoding='utf-8') as f:
        json.dump(formatted_xzfg, f, ensure_ascii=False)

    with open(formatted_dfxfg_path, 'w', encoding='utf-8') as f:
        json.dump(formatted_dfxfg, f, ensure_ascii=False)

    print(law_id)


def get_all_law_name_list():
    """
    This method may be abandoned.

    :return:
    """


    formatted_fl, formatted_sfjs, formatted_xzfg, formatted_dfxfg = read_formatted_laws()

    # FIXME only fl and sfjs
    all_formatted_laws = formatted_fl + formatted_sfjs

    all_law_name_list = [i['title'] for i in all_formatted_laws]
    all_law_name_list = list(set(all_law_name_list))

    with open(all_formatted_laws_name_path, 'w', encoding='utf-8') as f:
        json.dump(all_law_name_list, f, ensure_ascii=False)

    shorten_all_law_name_list = []
    for i in all_law_name_list:
        if i.startswith('中华人民共和国'):
            shorten_all_law_name_list.append(i[7:])
        elif i.startswith('最高人民法院、最高人民检察院'):
            shorten_all_law_name_list.append(i[14:])
        elif i.startswith('最高人民法院'):
            shorten_all_law_name_list.append(i[6:])
        elif i.startswith('最高人民检察院'):
            shorten_all_law_name_list.append(i[7:])

    shorten_all_law_name_list = list(set(shorten_all_law_name_list))
    with open(all_shorten_formatted_laws_name_path, 'w', encoding='utf-8') as f:
        json.dump(shorten_all_law_name_list, f, ensure_ascii=False)

    print(shorten_all_law_name_list)


def tmp():
    with open('/data/private/hxy/data/all_national_laws/formatted_laws/formatted_dfxfg.json', encoding='utf-8') as f:
        data = json.load(f)

    for d in data:
        print(d['title'])


if __name__ == '__main__':
    # format_pku_fl_sfjs_with_date(pku_fl_dir_path, TYPE_FL, fl_name_date_path, formatted_fl_path)
    # format_pku_fl_sfjs_with_date(pku_sfjs_dir_path, TYPE_SFJS, sfjs_name_date_path, formatted_sfjs_path)
    # format_national_xzfg_dfxfg(national_xzfg_dir_path, TYPE_XZFG, formatted_xzfg_path)
    # format_national_xzfg_dfxfg(national_dfxfg_dir_path, TYPE_DFXFG, formatted_dfxfg_path)
    # add_law_ids()
    # get_all_law_name_list()
